from pathlib import Path
from zipfile import ZipFile
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn

OUT = Path('downloads/参考案例2_测量水果电池的电动势和内阻_网络精选递进版.docx')
ASSET = Path('build_assets/fruit_battery')
OUT.parent.mkdir(parents=True, exist_ok=True)
ASSET.mkdir(parents=True, exist_ok=True)

SRC1 = 'https://zy.21cnjy.com/25750539'
SRC2 = 'https://tiku.baidu.com/tikupc/chapterdetail/ce16650e52ea551810a6870d-1561-15-jiaocai'
SRC3 = 'https://tiku.baidu.com/tikupc/chapterdetail/485364d97f1922791688e84d-116-5-knowpoint'


def set_font(run, east='宋体', latin='Times New Roman', size=12, bold=False, italic=False):
    run.font.name = latin
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run._element.rPr.rFonts.set(qn('w:eastAsia'), east)
    run._element.rPr.rFonts.set(qn('w:ascii'), latin)
    run._element.rPr.rFonts.set(qn('w:hAnsi'), latin)


def add_cn(p, text, size=12, bold=False):
    r = p.add_run(text)
    set_font(r, '宋体', 'Times New Roman', size, bold=bold)
    return r


def add_var(p, text, size=12):
    r = p.add_run(text)
    set_font(r, 'Times New Roman', 'Times New Roman', size, italic=True)
    return r


def add_roman(p, text, size=12, sub=False, sup=False):
    r = p.add_run(text)
    set_font(r, 'Times New Roman', 'Times New Roman', size)
    r.font.subscript = sub
    r.font.superscript = sup
    return r


def add_title(doc, text, size=18):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_font(r, '黑体', 'Times New Roman', size, bold=True)
    p.paragraph_format.space_after = Pt(5)


def add_heading(doc, text, page_break=False):
    if page_break:
        doc.add_page_break()
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r, '黑体', 'Times New Roman', 14, bold=True)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)


def add_label(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r, '黑体', 'Times New Roman', 12, bold=True)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(2)


def add_body(doc, text='', indent=True):
    p = doc.add_paragraph()
    if text:
        add_cn(p, text)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(3)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    return p


def add_source(doc, title, url):
    p = add_body(doc, indent=False)
    add_cn(p, title)
    add_cn(p, '；网络来源：')
    add_roman(p, url, size=9)


# ---------- OMML ----------
def mr(text, upright=False):
    style = 'p' if upright else 'i'
    return f'<m:r><m:rPr><m:sty m:val="{style}"/></m:rPr><w:rPr><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:eastAsia="Times New Roman"/></w:rPr><m:t>{text}</m:t></m:r>'


def mv(text): return mr(text, False)
def mu(text): return mr(text, True)
def msub(base, sub): return f'<m:sSub><m:sSubPr/><m:e>{base}</m:e><m:sub>{sub}</m:sub></m:sSub>'
def msup(base, sup): return f'<m:sSup><m:sSupPr/><m:e>{base}</m:e><m:sup>{sup}</m:sup></m:sSup>'
def mfrac(num, den): return f'<m:f><m:fPr/><m:num>{num}</m:num><m:den>{den}</m:den></m:f>'
def mdelim(expr): return f'<m:d><m:dPr/><m:e>{expr}</m:e></m:d>'


def add_formula(doc, expr):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    xml = f'<m:oMathPara {nsdecls("m", "w")}><m:oMath>{expr}</m:oMath></m:oMathPara>'
    p._element.append(parse_xml(xml))
    p.paragraph_format.space_after = Pt(4)


def E(): return mv('E')
def U(): return mv('U')
def I(): return mv('I')
def R(): return mv('R')
def r(): return mv('r')
def k(): return mv('k')
def b(): return mv('b')
def R0(): return msub(mv('R'), mu('0'))


# ---------- diagrams ----------
def font(sz):
    try:
        return ImageFont.truetype('/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf', sz)
    except Exception:
        return None


def draw_meter(path):
    img = Image.new('RGB', (900, 360), 'white')
    d = ImageDraw.Draw(img)
    f = font(28)
    d.arc((140, 40, 760, 330), 200, 340, fill='black', width=4)
    cx, cy, rad = 450, 300, 235
    import math
    for i in range(11):
        ang = math.radians(200 + 14*i)
        x1 = cx + (rad-18)*math.cos(ang); y1 = cy + (rad-18)*math.sin(ang)
        x2 = cx + rad*math.cos(ang); y2 = cy + rad*math.sin(ang)
        d.line((x1,y1,x2,y2), fill='black', width=3)
        if i % 2 == 0:
            val = 2.5*i/10
            tx = cx + (rad-55)*math.cos(ang)-16; ty = cy + (rad-55)*math.sin(ang)-14
            d.text((tx,ty), f'{val:.1f}', fill='black', font=f)
    ang = math.radians(200 + 14*7.2)
    px = cx + (rad-38)*math.cos(ang); py = cy + (rad-38)*math.sin(ang)
    d.line((cx,cy,px,py), fill='black', width=6)
    d.ellipse((438,288,462,312), fill='black')
    d.text((360,315), '2.5 V range', fill='black', font=f)
    img.save(path)


def draw_circuit(path, mode='basic'):
    img = Image.new('RGB', (1100, 540), 'white')
    d = ImageDraw.Draw(img); f=font(28); fs=font(23)
    # rectangle circuit
    d.line((120,130,890,130), fill='black', width=5)
    d.line((890,130,890,390), fill='black', width=5)
    d.line((890,390,120,390), fill='black', width=5)
    d.line((120,390,120,130), fill='black', width=5)
    # source
    d.line((120,235,120,275), fill='white', width=9)
    d.line((85,240,155,240), fill='black', width=4)
    d.line((98,275,142,275), fill='black', width=4)
    d.text((30,300), 'fruit cell', fill='black', font=fs)
    # ammeter
    d.ellipse((300,78,410,188), outline='black', width=4)
    d.text((340,112),'A',fill='black',font=f)
    # resistor
    d.rectangle((560,90,720,170), outline='black', width=4)
    d.text((615,108),'R',fill='black',font=f)
    # switch
    d.line((780,390,835,390), fill='white', width=8)
    d.line((770,390,810,350), fill='black', width=4)
    d.ellipse((765,383,779,397), fill='black'); d.ellipse((830,383,844,397), fill='black')
    d.text((800,410),'S',fill='black',font=fs)
    # voltmeter branch across R
    d.line((550,130,550,265), fill='black', width=4)
    d.line((730,130,730,265), fill='black', width=4)
    d.line((550,265,605,265), fill='black', width=4)
    d.line((675,265,730,265), fill='black', width=4)
    d.ellipse((605,220,675,290), outline='black', width=4)
    d.text((628,232),'V',fill='black',font=fs)
    img.save(path)


def draw_ui_graph(path):
    img = Image.new('RGB', (950, 620), 'white')
    d=ImageDraw.Draw(img); f=font(25); fs=font(21)
    ox, oy = 110, 520; xmax, ymax = 850, 80
    d.line((ox,oy,xmax,oy),fill='black',width=4)
    d.line((ox,oy,ox,ymax),fill='black',width=4)
    d.text((800,535),'I/mA',fill='black',font=f)
    d.text((25,65),'U/V',fill='black',font=f)
    d.text((55,95),'1.17',fill='black',font=fs)
    # lines same intercept, progressively smaller slopes a,b,c
    colors=['black','gray','black']
    ends=[(430,520),(600,520),(780,520)]
    labels=['a','b','c']
    for idx,(ex,ey) in enumerate(ends):
        d.line((ox,110,ex,ey), fill=colors[idx], width=4)
        mx=(ox+ex)//2; my=(110+ey)//2
        d.text((mx+10,my-30), labels[idx], fill='black', font=f)
    # point on b: 0.60 mA, 1.00 V, schematic scale
    px=ox+(600-ox)*0.60/(1.17/0.283333333)
    # easier direct approximate point on b
    px=210; py=165
    d.ellipse((px-6,py-6,px+6,py+6),fill='black')
    d.text((px+15,py-10),'(0.60 mA, 1.00 V)',fill='black',font=fs)
    img.save(path)


def draw_resistance_box_circuit(path):
    img=Image.new('RGB',(1050,500),'white')
    d=ImageDraw.Draw(img); f=font(28); fs=font(22)
    d.line((100,120,920,120),fill='black',width=5)
    d.line((920,120,920,360),fill='black',width=5)
    d.line((920,360,100,360),fill='black',width=5)
    d.line((100,360,100,120),fill='black',width=5)
    d.line((100,225,100,265),fill='white',width=9)
    d.line((65,230,135,230),fill='black',width=4); d.line((78,265,122,265),fill='black',width=4)
    d.text((20,285),'fruit cell',fill='black',font=fs)
    d.rectangle((320,80,480,160),outline='black',width=4); d.text((365,98),'R0',fill='black',font=f)
    d.rectangle((570,80,760,160),outline='black',width=4); d.text((625,98),'R box',fill='black',font=f)
    d.ellipse((765,70,875,180),outline='black',width=4); d.text((805,105),'A',fill='black',font=f)
    d.text((390,390),'multimeter used as ammeter',fill='black',font=fs)
    img.save(path)

meter=ASSET/'meter.png'; circuit=ASSET/'circuit.png'; graph=ASSET/'ui_graph.png'; boxc=ASSET/'box_circuit.png'
draw_meter(meter); draw_circuit(circuit); draw_ui_graph(graph); draw_resistance_box_circuit(boxc)

# ---------- document ----------
doc=Document()
sec=doc.sections[0]
sec.top_margin=Cm(2.0); sec.bottom_margin=Cm(2.0); sec.left_margin=Cm(2.2); sec.right_margin=Cm(2.2)
doc.styles['Normal'].font.name='Times New Roman'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
doc.styles['Normal'].font.size=Pt(12)

add_title(doc,'参考案例2  测量水果电池的电动势和内阻',18)
add_title(doc,'网络精选题组｜难度由低到高',14)
add_body(doc,'选题说明：三道题均来自公开网络题源，围绕水果电池内阻较大、仪器量程选择、U-I图像以及电阻箱法展开。电路图和坐标图按原题物理关系重新绘制，不改变核心设问。',indent=False)

# Q1
add_heading(doc,'一、基础题——粗测与器材选择')
add_label(doc,'【来源】')
add_source(doc,'21世纪教育网公开练习中的水果电池组实验题',SRC1)
add_label(doc,'【题目】')
add_body(doc,'寒假期间，某课外活动小组用苹果自制水果电池组。该电池组的电动势约为2 V，内阻在1 kΩ～2 kΩ之间。实验室有多用电表、电压表、滑动变阻器和开关、导线等器材。')
add_body(doc,'（1）用多用电表直流2.5 V挡粗测电动势，表针位置如图，读数是多少？')
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run().add_picture(str(meter),width=Cm(14.0))
add_body(doc,'（2）采用伏安法进一步测量，应选用量程0～3 V还是0～15 V的电压表？电流应选多用电表的0～1 mA挡还是0～100 mA挡？')
add_body(doc,'（3）滑动变阻器有R1（0～10 Ω）和R2（0～2000 Ω），应选择哪一个？')
add_label(doc,'【答案】')
add_body(doc,'（1）1.8 V；（2）电压表选0～3 V量程，多用电表选0～1 mA直流电流挡；（3）选R2。')
add_label(doc,'【解析】')
add_body(doc,'水果电池的电动势约2 V，电压表选0～3 V量程可保证读数精度；水果电池内阻为千欧量级，回路电流通常不超过毫安量级，因此应选0～1 mA挡。滑动变阻器的最大阻值应与水果电池内阻同一数量级，才能使电压和电流有明显变化，所以选择R2。')

# Q2
add_heading(doc,'二、提升题——电极插入深度与U-I图像',page_break=True)
add_label(doc,'【来源】')
add_source(doc,'百度题库公开的水果电池实验题',SRC2)
add_label(doc,'【题目】')
add_body(doc,'某同学用铜片和锌片制作水果电池，利用电压表、电流表、滑动变阻器、开关和导线测量其电动势和内阻。滑动变阻器有R1（0～20 Ω）和R2（0～2 kΩ）。')
add_body(doc,'（1）应选择哪一个滑动变阻器？')
add_body(doc,'（2）画出测量电路。')
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run().add_picture(str(circuit),width=Cm(15.0))
add_body(doc,'（3）保持铜片和锌片间距不变，逐渐增加插入深度，得到图中的a、b、c三条U-I图线。水果电池的电动势和内阻分别怎样变化？')
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run().add_picture(str(graph),width=Cm(14.5))
add_body(doc,'（4）图线b与电压轴的交点为1.17 V，图线上一点的坐标为I=0.60 mA、U=1.00 V，求图线b对应的内阻。')
add_label(doc,'【答案】')
add_body(doc,'（1）R2；（2）电流表串联，电压表测路端电压，滑动变阻器采用限流接法；（3）电动势基本不变，内阻减小；（4）约2.83×10² Ω。')
add_label(doc,'【解析】')
add_body(doc,'水果电池的内阻较大，应选择千欧量级的滑动变阻器。U-I图像与电压轴的交点表示电动势，三条图线纵截距相同，说明电动势基本不变；插入越深，短路电流越大，图线斜率绝对值越小，说明内阻减小。')
add_body(doc,'根据闭合电路欧姆定律可得')
add_formula(doc,U()+mu('=')+E()+mu('−')+I()+r())
add_body(doc,'因此')
add_formula(doc,r()+mu('=')+mfrac(E()+mu('−')+U(),I())+mu('=')+mfrac(mu('1.17−1.00'),mu('0.60×10'))+msup(mu('10'),mu('−3'))+mu(' Ω')+mu('≈283 Ω'))

# Q3
add_heading(doc,'三、综合题——多用电表与电阻箱法',page_break=True)
add_label(doc,'【来源】')
add_source(doc,'百度题库公开的苹果水果电池实验题',SRC3)
add_label(doc,'【题目】')
add_body(doc,'某中学生课外科技活动小组利用铜片、锌片和苹果制作水果电池。现有多用电表一个、高阻值电阻箱R一个、保护电阻R0一只、开关和导线若干。')
add_body(doc,'（1）锌片为负极、铜片为正极。用多用电表直流电压挡估测水果电池电动势时，红表笔应接铜片还是锌片？能否用多用电表欧姆挡估测水果电池的内阻？')
add_body(doc,'（2）把多用电表当作电流表使用，按图连接电路。改变电阻箱阻值R，测得多组电流I。若作1/I-R图像，图线斜率为k、纵截距为b，求水果电池的电动势E和内阻r。')
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run().add_picture(str(boxc),width=Cm(15.0))
add_label(doc,'【答案】')
add_body(doc,'（1）红表笔接铜片；不能用欧姆挡测量。')
p=add_body(doc,'（2）',indent=True)
add_var(p,'E'); add_roman(p,'='); add_roman(p,'1/'); add_var(p,'k'); add_cn(p,'，')
add_var(p,'r'); add_roman(p,'='); add_var(p,'b'); add_roman(p,'/'); add_var(p,'k'); add_roman(p,'−'); add_var(p,'R'); add_roman(p,'0',sub=True); add_cn(p,'。')
add_label(doc,'【解析】')
add_body(doc,'多用电表测直流电压时，电流应从红表笔流入，故红表笔接正极铜片。欧姆挡内部自带电源，而水果电池本身也有电动势，不能直接用欧姆挡测其内阻。')
add_body(doc,'根据闭合电路欧姆定律可得')
add_formula(doc,E()+mu('=')+I()+mdelim(R0()+mu('+')+R()+mu('+')+r()))
add_body(doc,'整理得')
add_formula(doc,mfrac(mu('1'),I())+mu('=')+mfrac(mu('1'),E())+R()+mu('+')+mfrac(R0()+mu('+')+r(),E()))
add_body(doc,'因此，1/I-R图像的斜率和截距分别满足')
add_formula(doc,k()+mu('=')+mfrac(mu('1'),E())+mu('，')+b()+mu('=')+mfrac(R0()+mu('+')+r(),E()))
add_body(doc,'所以')
add_formula(doc,E()+mu('=')+mfrac(mu('1'),k())+mu('，')+r()+mu('=')+mfrac(b(),k())+mu('−')+R0())

# self-check
add_heading(doc,'自检记录',page_break=True)
for item in [
    '三道题均来自公开网络题源，难度依次为器材选择、U-I图像分析、函数变换与参数提取。',
    '第2题由给定数据计算内阻：r=(1.17 V−1.00 V)/(0.60 mA)=283 Ω。',
    '第3题推导结果为E=1/k，r=b/k−R0。',
    '公式采用Word可编辑OMML结构，物理量斜体，数字、单位和说明性角标正体。',
    '文档未使用Unicode伪上标、伪下标，题目、答案和解析相互一致。',
]:
    add_body(doc,item)

doc.save(OUT)

assert OUT.exists() and OUT.stat().st_size > 25000
with ZipFile(OUT) as z:
    names=set(z.namelist())
    assert 'word/document.xml' in names and 'word/styles.xml' in names
    assert len([n for n in names if n.startswith('word/media/')]) >= 4
    xml=z.read('word/document.xml').decode('utf-8')
    forbidden='⁰¹²³⁴⁵⁶⁷⁸⁹⁻₀₁₂₃₄₅₆₇₈₉ₑₚ'
    assert not any(ch in xml for ch in forbidden)
    assert xml.count('【题目】')==3
    assert xml.count('【答案】')==3
    assert xml.count('【解析】')==3
    assert '<m:oMathPara' in xml and '<m:f>' in xml and '<m:sSub>' in xml and '<m:sSup>' in xml
    for text in ['1.8 V','R2','283 Ω','E=1/k','r=b/k−R0']:
        assert text in xml
    for bad in ['TODO','待补充','XXXX','�']:
        assert bad not in xml

print(OUT)
