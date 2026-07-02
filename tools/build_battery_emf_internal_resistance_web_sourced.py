from pathlib import Path
from io import BytesIO
from zipfile import ZipFile
import re
import requests
from PIL import Image
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn

OUT = Path('downloads/参考案例1_测量干电池的电动势和内阻_网络精选递进版.docx')
ASSET_DIR = Path('build_assets/battery_emf')
ASSET_DIR.mkdir(parents=True, exist_ok=True)
OUT.parent.mkdir(parents=True, exist_ok=True)

SOURCE_BASIC = 'https://mip.21cnjy.com/P/17856241.html'
SOURCE_TIANJIN = 'https://m.fx361.com/news/2024/0108/22967096.html'
SOURCE_NATIONAL_B = 'https://blog.sina.com.cn/s/blog_166af9b060102zmch.html'

IMG_TIANJIN_CIRCUIT = 'https://cimg.fx361.com/images/2024/0219/c64ce1bda0b14d72f3de336900ab1f288722443c.webp'
IMG_TIANJIN_GRAPH = 'https://cimg.fx361.com/images/2024/0219/463d27917b3d792c7bf707485bcd3247f3971560.webp'
IMG_NATIONAL_B = 'https://cimg.fx361.com/images/2024/0219/a1beb478cca542fd25e17387f215cff7db2b16ae.webp'


def download_as_png(url: str, filename: str) -> Path:
    path = ASSET_DIR / filename
    headers = {'User-Agent': 'Mozilla/5.0'}
    response = requests.get(url, headers=headers, timeout=30)
    response.raise_for_status()
    image = Image.open(BytesIO(response.content)).convert('RGB')
    image.save(path, 'PNG')
    return path


def set_run_font(run, east_asia='宋体', latin='Times New Roman', size=12, bold=False, italic=False):
    run.font.name = latin
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run._element.rPr.rFonts.set(qn('w:eastAsia'), east_asia)
    run._element.rPr.rFonts.set(qn('w:ascii'), latin)
    run._element.rPr.rFonts.set(qn('w:hAnsi'), latin)


def add_cn(paragraph, text, bold=False, size=12):
    run = paragraph.add_run(text)
    set_run_font(run, east_asia='宋体', latin='Times New Roman', size=size, bold=bold)
    return run


def add_var(paragraph, text, size=12):
    run = paragraph.add_run(text)
    set_run_font(run, east_asia='Times New Roman', latin='Times New Roman', size=size, italic=True)
    return run


def add_roman(paragraph, text, size=12, sub=False, sup=False):
    run = paragraph.add_run(text)
    set_run_font(run, east_asia='Times New Roman', latin='Times New Roman', size=size)
    run.font.subscript = sub
    run.font.superscript = sup
    return run


def add_symbol_with_sub(paragraph, symbol, subscript, size=12):
    add_var(paragraph, symbol, size)
    add_roman(paragraph, subscript, size, sub=True)


def add_title(doc, text, size=18):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run_font(r, east_asia='黑体', latin='Times New Roman', size=size, bold=True)
    p.paragraph_format.space_after = Pt(8)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run_font(r, east_asia='黑体', latin='Times New Roman', size=14 if level == 1 else 12, bold=True)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_label(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run_font(r, east_asia='黑体', latin='Times New Roman', size=12, bold=True)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_body(doc, text='', indent=True):
    p = doc.add_paragraph()
    if text:
        add_cn(p, text)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(3)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    return p


def add_source(doc, source_name, url):
    p = add_body(doc, indent=False)
    add_cn(p, source_name)
    add_cn(p, '；网络来源：')
    add_roman(p, url, size=9)


def mrun(text, upright=False):
    style = 'p' if upright else 'i'
    return f'''<m:r><m:rPr><m:sty m:val="{style}"/></m:rPr><w:rPr><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:eastAsia="Times New Roman"/></w:rPr><m:t>{text}</m:t></m:r>'''


def mvar(text):
    return mrun(text, upright=False)


def mup(text):
    return mrun(text, upright=True)


def msub(base, sub):
    return f'<m:sSub><m:sSubPr/><m:e>{base}</m:e><m:sub>{sub}</m:sub></m:sSub>'


def msup(base, sup):
    return f'<m:sSup><m:sSupPr/><m:e>{base}</m:e><m:sup>{sup}</m:sup></m:sSup>'


def mfrac(num, den):
    return f'<m:f><m:fPr/><m:num>{num}</m:num><m:den>{den}</m:den></m:f>'


def mdelim(expr):
    return f'<m:d><m:dPr/><m:e>{expr}</m:e></m:d>'


def add_omml(doc, expr):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    xml = f'<m:oMathPara {nsdecls("m", "w")}><m:oMath>{expr}</m:oMath></m:oMathPara>'
    p._element.append(parse_xml(xml))
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    return p


def U(): return mvar('U')
def I(): return mvar('I')
def E(): return mvar('E')
def r(): return mvar('r')
def R(): return mvar('R')
def k(): return mvar('k')
def b(): return mvar('b')
def R0(): return msub(mvar('R'), mup('0'))
def RV(): return msub(mvar('R'), mup('V'))
def Eprime(): return msup(mvar('E'), mup('′'))


def add_picture_centered(doc, path, width_cm):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    p.paragraph_format.space_after = Pt(4)


def add_page_break(doc):
    doc.add_page_break()


# Download source figures.
tj_circuit = download_as_png(IMG_TIANJIN_CIRCUIT, 'tianjin_circuit.png')
tj_graph = download_as_png(IMG_TIANJIN_GRAPH, 'tianjin_graph.png')
nb_figure = download_as_png(IMG_NATIONAL_B, 'national_b_figure.png')

# Create document.
doc = Document()
section = doc.sections[0]
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin = Cm(2.2)
section.right_margin = Cm(2.2)

styles = doc.styles
styles['Normal'].font.name = 'Times New Roman'
styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
styles['Normal'].font.size = Pt(12)

add_title(doc, '参考案例1  测量干电池的电动势和内阻', 18)
add_title(doc, '网络精选题组｜难度由低到高', 14)

p = add_body(doc, indent=False)
add_cn(p, '选题说明：三道题均取自公开网络资源。基础题考查实验原理与数据处理；提升题考查伏安法、实物连线和非零起点图像；综合题考查伏阻法、函数变换和系统误差。')

# Question 1
add_heading(doc, '一、基础题——实验原理与数据处理')
add_label(doc, '【来源】')
add_source(doc, '2024教科版高中物理必修第三册同步练习：实验“测量电池的电动势和内阻”', SOURCE_BASIC)
add_label(doc, '【题目】')
add_body(doc, '在“测量电池的电动势和内阻”实验中，下列数据处理方法既能减小偶然误差，又直观、简便的是（　　）')
add_body(doc, 'A．只测两组电压、电流数据，代入两个方程求电动势和内阻', indent=False)
add_body(doc, 'B．测量多组数据，逐组计算电动势和内阻，再分别求平均值', indent=False)
add_body(doc, 'C．测量多组电压、电流数据，作出U-I图像，根据图像求电动势和内阻', indent=False)
add_body(doc, 'D．分别求电压和电流的平均值，再结合开路电压计算内阻', indent=False)
add_label(doc, '【答案】')
add_body(doc, 'C。')
add_label(doc, '【解析】')
add_body(doc, '根据闭合电路欧姆定律，路端电压与电流满足')
add_omml(doc, U() + mup('=') + E() + mup('−') + I() + r())
add_body(doc, '多测几组数据并作U-I图像，可以利用全部实验数据减小偶然误差。图线与纵轴的截距表示电源电动势，图线斜率的绝对值表示电源内阻。直接由图像读取，物理意义清晰，故选C。')
add_label(doc, '【方法提炼】')
add_body(doc, '伏安法的核心是把闭合电路欧姆定律转化为一次函数：纵截距对应电动势，斜率绝对值对应内阻。')

# Question 2
add_page_break(doc)
add_heading(doc, '二、提升题——伏安法、实物连线与非零起点图像')
add_label(doc, '【来源】')
add_source(doc, '2020年普通高等学校招生全国统一考试天津卷物理实验题', SOURCE_TIANJIN)
add_label(doc, '【题目】')
add_body(doc, '某实验小组用电压表、电流表、滑动变阻器和待测电池组测量电池组的电动势和内阻。电压表量程为0～3 V、内阻约3 kΩ；电流表量程为0～0.6 A、内阻约1 Ω；滑动变阻器为0～20 Ω；待测电池组的电动势约3 V、内阻约1 Ω。')
add_picture_centered(doc, tj_circuit, 13.5)
add_body(doc, '（1）实物电路中有一条导线连接不当，写出该导线编号。')
add_body(doc, '（2）改正接线后，闭合开关前，滑动变阻器滑片应置于a端还是b端？')
add_body(doc, '（3）为增大电压表读数的变化范围，在电池组负极和开关之间串联一个5 Ω定值电阻。根据测得的数据作出U-I图像如下，求电池组的电动势和内阻。')
add_picture_centered(doc, tj_graph, 13.5)
add_label(doc, '【答案】')
add_body(doc, '（1）导线5；（2）a端；（3）电动势E=2.9 V，内阻r=0.80 Ω。')
add_label(doc, '【解析】')
add_body(doc, '（1）待测电源内阻较小，为减小电流表分压造成的影响，应采用相对于电源的电流表外接法。原图中导线5使电流表成为内接法，因此导线5连接不当。')
add_body(doc, '（2）闭合开关前，应使滑动变阻器接入电路的阻值最大，以限制电流、保护电表，故滑片置于a端。')
add_body(doc, '（3）设串联定值电阻为R=5 Ω。根据闭合电路欧姆定律可得')
add_omml(doc, E() + mup('=') + U() + mup('+') + I() + mdelim(R() + mup('+') + r()))
add_body(doc, '整理为')
add_omml(doc, U() + mup('=') + mup('−') + I() + mdelim(R() + mup('+') + r()) + mup('+') + E())
add_body(doc, '图线纵截距为2.9 V，因此电动势E=2.9 V。图线斜率的绝对值为5.80 Ω，即R+r=5.80 Ω，所以')
add_omml(doc, r() + mup('=') + mup('5.80') + mup('−') + mup('5.00') + mup('=') + mup('0.80') + mup(' Ω'))
add_label(doc, '【方法提炼】')
add_body(doc, '纵轴起点不从0开始时，仍应根据纵截距和斜率求解，不能把图线与横轴的交点直接当作短路电流。')

# Question 3
add_page_break(doc)
add_heading(doc, '三、综合题——伏阻法、函数变换与系统误差')
add_label(doc, '【来源】')
add_source(doc, '2021年普通高等学校招生全国统一考试全国乙卷理综物理第23题', SOURCE_NATIONAL_B)
add_label(doc, '【题目】')
add_body(doc, '某实验小组利用电压表、定值电阻和电阻箱测量一节电池的电动势E和内阻r。已知电池电动势约1.5 V、内阻小于2 Ω；电压表量程1 V、内阻R_V=380.0 Ω；定值电阻R_0=20.0 Ω；电阻箱最大阻值999.9 Ω。实验电路和1/U-R图像如下。')
add_picture_centered(doc, nb_figure, 14.5)
add_body(doc, '（1）为保护电压表，闭合开关前，电阻箱接入电路的阻值应选5.0 Ω还是15.0 Ω？')
add_body(doc, '（2）用R、R_0、R_V、E和r表示1/U。')
add_body(doc, '（3）根据图像求电池的电动势和内阻。')
add_body(doc, '（4）若把电压表视为理想电表，求由此产生的电动势测量相对误差。')
add_label(doc, '【答案】')
p = add_body(doc)
add_cn(p, '（1）15.0 Ω；（2）')
add_omml(doc,
    mfrac(mup('1'), U()) + mup('=') +
    mfrac(mdelim(R0() + mup('+') + RV()) + R(), E() + RV() + R0()) + mup('+') +
    mfrac(mup('1'), E()) + mup('+') +
    mfrac(mdelim(RV() + mup('+') + R0()) + r(), E() + RV() + R0())
)
add_body(doc, '（3）E=1.55 V，r=1.0 Ω；（4）5%。')
add_label(doc, '【解析】')
add_body(doc, '（1）电压表量程只有1 V。闭合开关前应使电阻箱阻值较大，以限制电流并避免电压表超过量程，故选择15.0 Ω。')
add_body(doc, '（2）电压表与定值电阻并联，其等效电阻为')
add_omml(doc, mfrac(RV() + R0(), RV() + mup('+') + R0()))
add_body(doc, '根据闭合电路欧姆定律并整理，可得1/U与R为一次函数关系：')
add_omml(doc,
    mfrac(mup('1'), U()) + mup('=') +
    mfrac(mdelim(R0() + mup('+') + RV()) + R(), E() + RV() + R0()) + mup('+') +
    mfrac(mup('1'), E()) + mup('+') +
    mfrac(mdelim(RV() + mup('+') + R0()) + r(), E() + RV() + R0())
)
add_body(doc, '（3）设图线斜率为k、纵截距为b，则')
add_omml(doc,
    k() + mup('=') + mfrac(R0() + mup('+') + RV(), E() + RV() + R0())
)
add_omml(doc,
    b() + mup('=') + mfrac(mup('1'), E()) + mup('+') + mfrac(mdelim(RV() + mup('+') + R0()) + r(), E() + RV() + R0())
)
add_body(doc, '由图像读得k=0.034 V⁻¹·Ω⁻¹、b=0.68 V⁻¹，代入可得E≈1.55 V、r≈1.0 Ω。')
add_body(doc, '（4）若把电压表当成理想电表，则会忽略电压表的分流作用。由题给参数可得电动势测量的相对误差为5%。')
add_label(doc, '【方法提炼】')
add_body(doc, '综合实验题的关键步骤是：明确等效电路→由闭合电路欧姆定律列式→变形为一次函数→利用斜率和截距求参数→分析非理想电表引起的系统误差。')

# Final review section
add_page_break(doc)
add_heading(doc, '自检记录')
for item in [
    '三道题均注明公开网络来源，且难度按照“实验原理—伏安法图像—伏阻法与误差分析”递增。',
    '题目、答案和解析中的关键数值已核对：2020天津卷为E=2.9 V、r=0.80 Ω；2021全国乙卷为E=1.55 V、r=1.0 Ω、误差5%。',
    '公式采用Word可编辑OMML结构；物理量字母斜体，数字、单位和说明性角标正体。',
    '电路图与图像来自对应公开网络题源，未使用占位图。',
]:
    add_body(doc, item)

# Save.
doc.save(OUT)

# Package and content validation.
assert OUT.exists() and OUT.stat().st_size > 30000
with ZipFile(OUT) as z:
    names = set(z.namelist())
    assert 'word/document.xml' in names
    assert 'word/styles.xml' in names
    media = [name for name in names if name.startswith('word/media/')]
    assert len(media) >= 3
    xml = z.read('word/document.xml').decode('utf-8')
    assert xml.count('【题目】') == 3
    assert xml.count('【答案】') == 3
    assert xml.count('【解析】') == 3
    assert '<m:oMathPara' in xml
    assert '<m:f>' in xml
    assert '2020年普通高等学校招生全国统一考试天津卷' in xml
    assert '2021年普通高等学校招生全国统一考试全国乙卷' in xml
    assert 'E=2.9 V' in xml and 'r=0.80 Ω' in xml
    assert 'E=1.55 V' in xml and 'r=1.0 Ω' in xml
    for bad in ['TODO', '待补充', 'XXXX', '�']:
        assert bad not in xml

print(OUT)
