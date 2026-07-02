from pathlib import Path
from zipfile import ZipFile
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn

OUT = Path('downloads/一题三变加高考真题_第十二章_电路中的能量转化_精修版.docx')
ASSET_DIR = Path('build_assets/circuit_energy_motor')
OUT.parent.mkdir(parents=True, exist_ok=True)
ASSET_DIR.mkdir(parents=True, exist_ok=True)

SOURCE_CLASSIC = 'https://www.maitianclub.com/uploadfile/attach/0/0/192.pdf'
SOURCE_LIFT = 'https://www.jyeoo.com/shiti/6410d174-e115-1556-9551-c225d03551df'
SOURCE_GAOKAO = 'https://zy.21cnjy.com/23525531'
SOURCE_GAOKAO_REVIEW = 'https://news.hangzhou.com.cn/zjnews/content/2022-01/10/content_8140110_3.htm'


def set_font(run, east_asia='宋体', latin='Times New Roman', size=12, bold=False, italic=False):
    run.font.name = latin
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run._element.rPr.rFonts.set(qn('w:eastAsia'), east_asia)
    run._element.rPr.rFonts.set(qn('w:ascii'), latin)
    run._element.rPr.rFonts.set(qn('w:hAnsi'), latin)


def add_cn(p, text, bold=False, size=12):
    r = p.add_run(text)
    set_font(r, east_asia='宋体', latin='Times New Roman', size=size, bold=bold)
    return r


def add_roman(p, text, size=12, sub=False, sup=False):
    r = p.add_run(text)
    set_font(r, east_asia='Times New Roman', latin='Times New Roman', size=size)
    r.font.subscript = sub
    r.font.superscript = sup
    return r


def add_var(p, text, size=12):
    r = p.add_run(text)
    set_font(r, east_asia='Times New Roman', latin='Times New Roman', size=size, italic=True)
    return r


def add_symbol_sub(p, symbol, subscript, size=12):
    add_var(p, symbol, size)
    add_roman(p, subscript, size, sub=True)


def add_title(doc, text, size=18):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_font(r, east_asia='黑体', latin='Times New Roman', size=size, bold=True)
    p.paragraph_format.space_after = Pt(5)
    return p


def add_heading(doc, text, page_break=False):
    if page_break:
        doc.add_page_break()
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r, east_asia='黑体', latin='Times New Roman', size=14, bold=True)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_label(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r, east_asia='黑体', latin='Times New Roman', size=12, bold=True)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_body(doc, text='', indent=True):
    p = doc.add_paragraph()
    if text:
        add_cn(p, text)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(3)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    return p


def add_source(doc, text, url):
    p = add_body(doc, indent=False)
    add_cn(p, text)
    add_cn(p, '；网络来源：')
    add_roman(p, url, size=9)


# ---------- OMML ----------
def mr(text, upright=False):
    style = 'p' if upright else 'i'
    return f'<m:r><m:rPr><m:sty m:val="{style}"/></m:rPr><w:rPr><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:eastAsia="Times New Roman"/></w:rPr><m:t>{text}</m:t></m:r>'


def mv(text): return mr(text, upright=False)
def mu(text): return mr(text, upright=True)
def msub(base, sub): return f'<m:sSub><m:sSubPr/><m:e>{base}</m:e><m:sub>{sub}</m:sub></m:sSub>'
def msup(base, sup): return f'<m:sSup><m:sSupPr/><m:e>{base}</m:e><m:sup>{sup}</m:sup></m:sSup>'
def mfrac(num, den): return f'<m:f><m:fPr/><m:num>{num}</m:num><m:den>{den}</m:den></m:f>'
def mdelim(expr): return f'<m:d><m:dPr/><m:e>{expr}</m:e></m:d>'


def add_formula(doc, expr):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    xml = f'<m:oMathPara {nsdecls("m", "w")}><m:oMath>{expr}</m:oMath></m:oMathPara>'
    p._element.append(parse_xml(xml))
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    return p


def P(sub): return msub(mv('P'), mu(sub))
def W(sub): return msub(mv('W'), mu(sub))
def E(): return mv('E')
def I(): return mv('I')
def U(): return mv('U')
def R(): return mv('R')
def r(): return mv('r')
def eta(): return mv('η')
def v(): return mv('v')
def m(): return mv('m')
def g(): return mv('g')
def H(): return mv('H')
def mdot(): return msub(mv('m'), mu('t'))


# ---------- Simple diagram ----------
def draw_lift_diagram(path: Path):
    img = Image.new('RGB', (1100, 500), 'white')
    d = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype('/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf', 32)
        small = ImageFont.truetype('/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf', 25)
    except Exception:
        font = None
        small = None
    # circuit line
    d.line((80, 130, 300, 130), fill='black', width=5)
    d.rectangle((300, 85, 470, 175), outline='black', width=5)
    d.text((330, 105), 'R=10 Ω', fill='black', font=small)
    d.line((470, 130, 610, 130), fill='black', width=5)
    d.ellipse((610, 70, 780, 240), outline='black', width=5)
    d.text((655, 120), 'M', fill='black', font=font)
    d.line((780, 130, 990, 130), fill='black', width=5)
    d.line((990, 130, 990, 350), fill='black', width=5)
    d.line((990, 350, 80, 350), fill='black', width=5)
    d.line((80, 350, 80, 130), fill='black', width=5)
    # source
    d.line((80, 230, 80, 265), fill='white', width=8)
    d.line((48, 235, 112, 235), fill='black', width=4)
    d.line((60, 265, 100, 265), fill='black', width=4)
    d.text((15, 285), '220 V', fill='black', font=small)
    # lift
    d.line((695, 240, 695, 405), fill='black', width=4)
    d.rectangle((620, 405, 770, 475), outline='black', width=4)
    d.text((650, 420), '10 kg', fill='black', font=small)
    d.text((800, 205), 'U_M=170 V', fill='black', font=small)
    img.save(path)


diagram_path = ASSET_DIR / 'motor_lift_diagram.png'
draw_lift_diagram(diagram_path)

# ---------- Document ----------
doc = Document()
sec = doc.sections[0]
sec.top_margin = Cm(2.0)
sec.bottom_margin = Cm(2.0)
sec.left_margin = Cm(2.2)
sec.right_margin = Cm(2.2)

doc.styles['Normal'].font.name = 'Times New Roman'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
doc.styles['Normal'].font.size = Pt(12)

add_title(doc, '第十二章  电能  能量守恒定律', 18)
add_title(doc, '1. 电路中的能量转化', 16)
add_title(doc, '课本例题·一题三变＋高考真题', 14)

# Original example
add_heading(doc, '课本例题原型')
add_label(doc, '【题目】')
add_body(doc, '一台电动机，线圈的电阻为6 Ω。当它两端所加的电压为220 V时，通过的电流为5 A。这台电动机发热的功率与对外做功的功率各是多少？')
add_label(doc, '【答案】')
add_body(doc, '发热功率为150 W；对外做功的功率为950 W。')
add_label(doc, '【解析】')
add_body(doc, '根据焦耳定律可得电动机线圈的发热功率')
add_formula(doc, P('热') + mu('=') + msup(I(), mu('2')) + R() + mu('=') + mdelim(msup(mu('5'), mu('2')) + mu('×6')) + mu(' W') + mu('=150 W'))
add_body(doc, '根据电功率公式可得电动机的输入功率')
add_formula(doc, P('电') + mu('=') + U() + I() + mu('=') + mdelim(mu('220×5')) + mu(' W') + mu('=1100 W'))
add_body(doc, '根据能量守恒定律可得电动机对外做功的功率')
add_formula(doc, P('机') + mu('=') + P('电') + mu('−') + P('热') + mu('=') + mdelim(mu('1100−150')) + mu(' W') + mu('=950 W'))
add_body(doc, '注意：电动机正常转动时是非纯电阻用电器，不能把电动机的总功率按纯电阻公式计算。')

# Basic variation
add_heading(doc, '一、基础变式——功率与能量的直接计算', page_break=True)
add_label(doc, '【题目】')
add_body(doc, '一台直流电动机正常工作时，两端电压为120 V，通过线圈的电流为4.0 A，线圈电阻为5.0 Ω。忽略除线圈发热外的其他能量损失。')
add_body(doc, '（1）求电动机的输入功率、发热功率和机械功率。')
add_body(doc, '（2）求电动机的效率。')
add_body(doc, '（3）电动机连续工作2.0 min，输入的电能、产生的热量和对外所做的机械功分别是多少？')
add_label(doc, '【答案】')
add_body(doc, '（1）480 W、80 W、400 W；（2）83.3%；（3）57.6 kJ、9.6 kJ、48.0 kJ。')
add_label(doc, '【解析】')
add_body(doc, '根据电功率公式可得输入功率')
add_formula(doc, P('电') + mu('=') + U() + I() + mu('=') + mdelim(mu('120×4.0')) + mu(' W') + mu('=480 W'))
add_body(doc, '根据焦耳定律可得发热功率')
add_formula(doc, P('热') + mu('=') + msup(I(), mu('2')) + R() + mu('=') + mdelim(msup(mu('4.0'), mu('2')) + mu('×5.0')) + mu(' W') + mu('=80 W'))
add_body(doc, '根据能量守恒定律可得机械功率')
add_formula(doc, P('机') + mu('=') + P('电') + mu('−') + P('热') + mu('=400 W'))
add_formula(doc, eta() + mu('=') + mfrac(P('机'), P('电')) + mu('×100%') + mu('=83.3%'))
add_body(doc, '根据功率与时间的关系，工作时间为120 s，可得')
add_formula(doc, W('电') + mu('=') + P('电') + mu('t') + mu('=57.6 kJ，') + W('热') + mu('=9.6 kJ，') + W('机') + mu('=48.0 kJ'))

# Context variation
add_heading(doc, '二、情境变式——转子静止、正常工作与堵转', page_break=True)
add_label(doc, '【来源】')
add_source(doc, '依据公开资料中的经典电动机例题整理', SOURCE_CLASSIC)
add_label(doc, '【题目】')
add_body(doc, '某直流电动机转子被固定不动时，在其两端加0.20 V电压，测得电流为0.40 A。转子能够正常转动后，在其两端加2.0 V电压，工作电流为1.0 A。忽略除线圈发热外的其他损耗。')
add_body(doc, '（1）求电动机线圈的电阻。')
add_body(doc, '（2）求电动机正常工作时的输入功率、发热功率、机械功率和效率。')
add_body(doc, '（3）若电动机在2.0 V电压下突然堵转，求此时的电流和发热功率。')
add_body(doc, '（4）说明为什么电动机正常工作时不能直接用欧姆定律求电流。')
add_label(doc, '【答案】')
add_body(doc, '（1）0.50 Ω；（2）2.0 W、0.50 W、1.5 W、75%；（3）4.0 A、8.0 W；（4）正常转动时电能还转化为机械能，电动机不是纯电阻元件。')
add_label(doc, '【解析】')
add_body(doc, '转子静止时没有机械能输出，电动机可视为纯电阻元件。根据欧姆定律可得')
add_formula(doc, R() + mu('=') + mfrac(mu('0.20 V'), mu('0.40 A')) + mu('=0.50 Ω'))
add_body(doc, '正常工作时，根据电功率公式和焦耳定律可得')
add_formula(doc, P('电') + mu('=') + U() + I() + mu('=2.0 W'))
add_formula(doc, P('热') + mu('=') + msup(I(), mu('2')) + R() + mu('=0.50 W'))
add_formula(doc, P('机') + mu('=') + P('电') + mu('−') + P('热') + mu('=1.5 W'))
add_formula(doc, eta() + mu('=') + mfrac(mu('1.5'), mu('2.0')) + mu('×100%=75%'))
add_body(doc, '堵转后电动机不再输出机械功率，重新成为纯电阻元件。')
add_formula(doc, I() + mu('=') + mfrac(mu('2.0 V'), mu('0.50 Ω')) + mu('=4.0 A'))
add_formula(doc, P('热') + mu('=') + msup(I(), mu('2')) + R() + mu('=8.0 W'))
add_body(doc, '堵转时电流显著增大，线圈发热功率急剧增大，因此电动机长时间堵转容易烧毁。')

# Literacy enhancement
add_heading(doc, '三、素养提升——提升重物与故障分析', page_break=True)
add_label(doc, '【来源】')
add_source(doc, '依据公开题库中的直流电动机提升重物问题整理', SOURCE_LIFT)
add_label(doc, '【题目】')
p = add_body(doc, '', indent=True)
add_cn(p, '如图，电源电压恒为220 V，保护电阻R=10 Ω，电动机线圈电阻r=1.0 Ω。电动机正常工作时，电压表测得电动机两端电压为170 V。电动机匀速提升质量为10 kg的重物，不计摩擦，g取')
add_roman(p, '10 m/s')
add_roman(p, '2', sup=True)
add_cn(p, '。')
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run().add_picture(str(diagram_path), width=Cm(15.0))
add_body(doc, '（1）求电动机正常工作时的电流。')
add_body(doc, '（2）求电动机的输入功率、线圈发热功率和机械功率。')
add_body(doc, '（3）求重物匀速上升的速度。')
add_body(doc, '（4）若电动机突然卡住，求电路中的电流以及电动机的输入功率和发热功率。')
add_label(doc, '【答案】')
add_body(doc, '（1）5.0 A；（2）850 W、25 W、825 W；（3）8.25 m/s；（4）20 A、400 W、400 W。')
add_label(doc, '【解析】')
add_body(doc, '正常工作时，保护电阻两端电压为220 V−170 V=50 V。根据欧姆定律可得电路电流')
add_formula(doc, I() + mu('=') + mfrac(mu('50 V'), mu('10 Ω')) + mu('=5.0 A'))
add_body(doc, '根据电功率公式可得电动机输入功率')
add_formula(doc, P('电') + mu('=') + mdelim(mu('170×5.0')) + mu(' W') + mu('=850 W'))
add_body(doc, '根据焦耳定律可得线圈发热功率')
add_formula(doc, P('热') + mu('=') + mdelim(msup(mu('5.0'), mu('2')) + mu('×1.0')) + mu(' W') + mu('=25 W'))
add_body(doc, '根据能量守恒定律可得机械功率')
add_formula(doc, P('机') + mu('=') + mu('850 W−25 W=825 W'))
add_body(doc, '重物匀速上升时，电动机的机械功率等于克服重力做功的功率。')
add_formula(doc, P('机') + mu('=') + m() + g() + v())
add_formula(doc, v() + mu('=') + mfrac(mu('825 W'), mdelim(mu('10 kg×10 m/') + msup(mu('s'), mu('2')))) + mu('=8.25 m/s'))
add_body(doc, '电动机卡住后，电动机只剩线圈电阻，电路总电阻为11 Ω。')
add_formula(doc, I() + mu('=') + mfrac(mu('220 V'), mu('11 Ω')) + mu('=20 A'))
add_formula(doc, P('电') + mu('=') + P('热') + mu('=') + msup(I(), mu('2')) + r() + mu('=400 W'))

# Gaokao
add_heading(doc, '四、高考真题——电动机与水泵的能量转化', page_break=True)
add_label(doc, '【真题来源】')
add_source(doc, '2022年1月浙江省普通高校招生选考科目考试物理第12题', SOURCE_GAOKAO)
add_label(doc, '【题目】')
p = add_body(doc, '', indent=True)
add_cn(p, '某节水喷灌系统中，水以15 m/s的速度水平喷出，每秒喷出水的质量为2.0 kg。喷出的水从井下抽取，喷口离水面的高度保持3.75 m不变。水泵由电动机带动，电动机正常工作时输入电压为220 V，输入电流为2.0 A。不计电动机的摩擦损耗，电动机的输出功率等于水泵的输入功率。水泵效率为75%，忽略水在管道中运动的机械能损失，g取')
add_roman(p, '10 m/s')
add_roman(p, '2', sup=True)
add_cn(p, '。下列说法正确的是（　　）')
add_body(doc, 'A．每秒水泵对水做功为75 J', indent=False)
add_body(doc, 'B．每秒水泵对水做功为225 J', indent=False)
add_body(doc, 'C．水泵输入功率为440 W', indent=False)
add_body(doc, 'D．电动机线圈的电阻为10 Ω', indent=False)
add_label(doc, '【答案】')
add_body(doc, 'D。')
add_label(doc, '【解析】')
add_body(doc, '每秒喷出的水获得的机械能包括重力势能和动能。根据机械能的增加率可得水泵对水的输出功率')
add_formula(doc, P('水') + mu('=') + mdot() + mdelim(g() + H() + mu('+') + mfrac(msup(v(), mu('2')), mu('2'))) + mu('=300 W'))
add_body(doc, '因此A、B错误。根据水泵效率可得水泵输入功率')
add_formula(doc, P('泵入') + mu('=') + mfrac(mu('300 W'), mu('75%')) + mu('=400 W'))
add_body(doc, '因此C错误。电动机输入功率为')
add_formula(doc, P('电') + mu('=') + mdelim(mu('220×2.0')) + mu(' W') + mu('=440 W'))
add_body(doc, '电动机的输出功率等于水泵输入功率，所以电动机线圈的发热功率为')
add_formula(doc, P('热') + mu('=') + mu('440 W−400 W=40 W'))
add_body(doc, '根据焦耳定律可得电动机线圈电阻')
add_formula(doc, R() + mu('=') + mfrac(P('热'), msup(I(), mu('2'))) + mu('=') + mfrac(mu('40 W'), msup(mu('2.0 A'), mu('2'))) + mu('=10 Ω'))
add_body(doc, '故选D。该题以真实节水喷灌系统为情境，综合考查机械能、电功率、效率和焦耳热。浙江省相关命题评析也将第12题列为利用真实情境考查能量观的代表题。')

# Self-check
add_heading(doc, '自检记录', page_break=True)
for item in [
    '原例题、三道递进题和高考真题均围绕“电动机输入功率—发热功率—机械功率”的能量守恒关系展开。',
    '难度递进为：直接计算→工作状态变化→电路、机械与故障综合→真实高考情境综合。',
    '关键数值已复核：原题150 W、950 W；基础题480 W、80 W、400 W；堵转题4.0 A、8.0 W；提升题5.0 A、825 W、8.25 m/s、20 A；高考题正确选项D。',
    '公式使用Word可编辑OMML结构；物理量为斜体，数字、单位和说明性角标为正体。',
    '未使用Unicode伪上标、伪下标；题目、答案和解析相互一致。',
]:
    add_body(doc, item)

# Save and validate
doc.save(OUT)
assert OUT.exists() and OUT.stat().st_size > 25000
with ZipFile(OUT) as z:
    names = set(z.namelist())
    assert 'word/document.xml' in names and 'word/styles.xml' in names
    assert any(n.startswith('word/media/') for n in names)
    xml = z.read('word/document.xml').decode('utf-8')
    forbidden = '⁰¹²³⁴⁵⁶⁷⁸⁹⁻₀₁₂₃₄₅₆₇₈₉ₑₚ'
    assert not any(ch in xml for ch in forbidden)
    assert xml.count('【题目】') == 5
    assert xml.count('【答案】') == 5
    assert xml.count('【解析】') == 5
    assert '<m:oMathPara' in xml and '<m:f>' in xml and '<m:sSub>' in xml and '<m:sSup>' in xml
    assert '2022年1月浙江省普通高校招生选考科目考试物理第12题' in xml
    assert '8.25 m/s' in xml and '20 A' in xml and 'D。' in xml
    for bad in ['TODO', '待补充', 'XXXX', '�']:
        assert bad not in xml

print(OUT)
